#!/usr/bin/env python3
"""Generate a programmer resume DOCX from normalized facts."""

from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BLUE = "0070C0"
GOLD = "C8A96A"
TEXT = "111111"
GRAY = "666666"
CANONICAL_TEMPLATE_NAME = "xxx.doc"
CANONICAL_TEMPLATE_HINT = "优秀简历汇总"
STYLE_SOURCE = "canonical xxx.doc template"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def select_template(index: dict) -> dict:
    files = index.get("files", [])
    canonical = [
        item
        for item in files
        if item.get("path", "").replace("\\", "/").endswith(f"{CANONICAL_TEMPLATE_HINT}/{CANONICAL_TEMPLATE_NAME}")
    ]
    if canonical:
        return canonical[0]
    named = [item for item in files if Path(item.get("path", "")).name == CANONICAL_TEMPLATE_NAME]
    return named[0] if named else (files[0] if files else {"path": f"{CANONICAL_TEMPLATE_HINT}/{CANONICAL_TEMPLATE_NAME}"})


def set_run_font(run, size: int = 10, bold: bool = False, color: str = TEXT, font: str = "Microsoft YaHei") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, **borders: dict) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, attrs in borders.items():
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def set_table_borderless(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
                insideH={"val": "nil"},
                insideV={"val": "nil"},
            )


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph._element.getparent().remove(paragraph._element)


def add_cell_text(cell, text: str, size: int = 10, bold: bool = False, color: str = TEXT, align=None) -> None:
    clear_cell(cell)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)


def compact_join(parts: Iterable[str], separator: str = "    ") -> str:
    return separator.join(part for part in parts if part)


def date_range(item: dict) -> str:
    start = item.get("start", "")
    end = item.get("end", "")
    if start and end:
        return f"{start}-{end}"
    return start or end


def meaningful_items(items: list[dict]) -> list[dict]:
    return [
        item
        for item in items
        if any(value for value in item.values() if isinstance(value, str))
        or any(item.get(key) for key in ["bullets", "technologies"])
    ]


def first_education(facts: dict) -> dict:
    education = meaningful_items(facts.get("education", []))
    return education[0] if education else {}


def flatten_skills(skills: dict) -> list[dict]:
    labels = {
        "languages": "编程语言",
        "backend": "后端框架",
        "frontend": "前端技术",
        "databases": "数据库",
        "tools": "工具平台",
    }
    rows = []
    for key, label in labels.items():
        values = skills.get(key) or []
        if values:
            rows.append({"label": label, "value": "、".join(values)})
    return rows


def add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(title)
    set_run_font(run, size=16, bold=True, color="000000", font="Arial")


def add_top_bar(document: Document) -> None:
    table = document.add_table(rows=1, cols=2)
    set_table_borderless(table)
    table.allow_autofit = False
    set_cell_width(table.cell(0, 0), 10.4)
    set_cell_width(table.cell(0, 1), 7.1)
    for index, color in enumerate([BLUE, GOLD]):
        cell = table.cell(0, index)
        set_cell_shading(cell, color)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_cell_text(cell, "", size=1)
    for cell in table.rows[0].cells:
        cell.height = Cm(0.22)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def add_section_heading(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=2)
    set_table_borderless(table)
    table.allow_autofit = False
    label_cell = table.cell(0, 0)
    line_cell = table.cell(0, 1)
    set_cell_width(label_cell, 4.0)
    set_cell_width(line_cell, 13.5)
    set_cell_shading(label_cell, BLUE)
    add_cell_text(label_cell, text, size=11, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(line_cell, "", size=1)
    set_cell_border(line_cell, bottom={"val": "single", "sz": "10", "space": "0", "color": BLUE})
    table.rows[0].height = Cm(0.55)


def add_basic_info(document: Document, rows: list[tuple[str, str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=0, cols=2)
    set_table_borderless(table)
    table.allow_autofit = False
    for left, right in zip(rows[0::2], rows[1::2]):
        row = table.add_row()
        for cell, item in zip(row.cells, [left, right]):
            set_cell_width(cell, 8.6)
            add_cell_text(cell, f"{item[0]}：{item[1]}", size=10)
    if len(rows) % 2:
        row = table.add_row()
        set_cell_width(row.cells[0], 8.6)
        set_cell_width(row.cells[1], 8.6)
        add_cell_text(row.cells[0], f"{rows[-1][0]}：{rows[-1][1]}", size=10)
        add_cell_text(row.cells[1], "", size=10)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_arrow_bullets(document: Document, lines: list[str], size: int = 9) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.35)
        paragraph.paragraph_format.space_after = Pt(2)
        marker = paragraph.add_run("➢  ")
        set_run_font(marker, size=size, bold=True, color="000000")
        run = paragraph.add_run(line)
        set_run_font(run, size=size, color="000000")


def add_project(document: Document, item: dict) -> None:
    heading = compact_join([item.get("period", ""), item.get("name", ""), item.get("role", "")], "    ")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(heading)
    set_run_font(run, size=10, bold=True)
    technologies = item.get("technologies", [])
    lines = []
    if technologies:
        lines.append(f"技术栈：{'、'.join(technologies)}")
    lines.extend(item.get("bullets", []))
    add_arrow_bullets(document, lines, size=8)


def add_work(document: Document, item: dict) -> None:
    heading = compact_join([date_range(item), item.get("company", ""), item.get("role", "")], "    ")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(heading)
    set_run_font(run, size=10, bold=True)
    add_arrow_bullets(document, item.get("bullets", []), size=8)


def render_plain_text(facts: dict, selected_template: dict) -> str:
    personal = facts.get("personal", {})
    contact_line = " | ".join(filter(None, [personal.get("phone"), personal.get("email"), personal.get("city"), personal.get("github")]))
    lines = [
        personal.get("name", ""),
        f"求职意向：{facts.get('target_role', '')}",
    ]
    if contact_line:
        lines.append(contact_line)
    lines.extend(["", f"模板来源：{selected_template.get('path', 'template corpus')}"])
    if facts.get("summary"):
        lines.extend(["", "个人优势", facts.get("summary", "")])
    skill_rows = flatten_skills(facts.get("skills", {}))
    if skill_rows:
        lines.extend(["", "开发技能"])
        for row in skill_rows:
            lines.append(f"{row['label']}：{row['value']}")
    for section, key in [("工作经历", "work_experience"), ("项目经历", "projects"), ("教育背景", "education")]:
        items = meaningful_items(facts.get(key, []))
        if not items:
            continue
        lines.extend(["", section])
        for item in items:
            if key == "work_experience":
                lines.append(compact_join([date_range(item), item.get("company", ""), item.get("role", "")], " "))
            elif key == "projects":
                lines.append(compact_join([item.get("period", ""), item.get("name", ""), item.get("role", "")], " "))
            elif key == "education":
                lines.append(compact_join([date_range(item), item.get("school", ""), item.get("major", ""), item.get("degree", "")], " "))
            for bullet in item.get("bullets", []):
                lines.append(f"- {bullet}")
    return "\n".join(line for line in lines if line is not None)


def build_render_payload(facts: dict, selected_template: dict) -> dict:
    personal = facts.get("personal", {})
    education = first_education(facts)
    basics = []
    for label, value in [
        ("姓名", personal.get("name")),
        ("联系电话", personal.get("phone")),
        ("电子邮件", personal.get("email")),
        ("所在城市", personal.get("city")),
        ("GitHub", personal.get("github")),
        ("专业背景", education.get("major")),
        ("毕业院校", education.get("school")),
        ("学历", education.get("degree")),
    ]:
        if value:
            basics.append((label, value))
    target_role = facts.get("target_role", "")
    intent = [("目标职位", target_role)] if target_role else []
    skill_rows = flatten_skills(facts.get("skills", {}))
    title = f"{target_role}_{personal.get('name', '')}".strip("_") or personal.get("name", "resume")
    return {
        "title": title,
        "personal": personal,
        "summary": facts.get("summary", ""),
        "basics": basics,
        "intent": intent,
        "skills": skill_rows,
        "work_experience": meaningful_items(facts.get("work_experience", [])),
        "projects": meaningful_items(facts.get("projects", [])),
        "education": meaningful_items(facts.get("education", [])),
        "certificates": facts.get("certificates", []),
        "selected_template": selected_template,
        "style": {
            "source": STYLE_SOURCE,
            "blue": BLUE,
            "gold": GOLD,
            "template_name": CANONICAL_TEMPLATE_NAME,
            "template_hint": CANONICAL_TEMPLATE_HINT,
        },
    }


def render_resume(facts: dict, selected_template: dict, out_dir: Path) -> Path:
    payload = build_render_payload(facts, selected_template)
    document = Document()
    configure_document(document)
    add_title(document, payload["title"])
    add_top_bar(document)

    if payload["basics"]:
        add_section_heading(document, "基本信息")
        add_basic_info(document, payload["basics"])

    if payload["intent"]:
        add_section_heading(document, "求职意向")
        add_basic_info(document, payload["intent"])

    if payload["skills"]:
        add_section_heading(document, "开发技能")
        add_arrow_bullets(document, [f"{row['label']}：{row['value']}" for row in payload["skills"]], size=9)

    if payload["work_experience"]:
        add_section_heading(document, "工作经历")
        for item in payload["work_experience"]:
            add_work(document, item)

    if payload["projects"]:
        add_section_heading(document, "项目经历")
        for item in payload["projects"]:
            add_project(document, item)

    if payload["education"] and not payload["basics"]:
        add_section_heading(document, "教育背景")
        education_rows: list[tuple[str, str]] = []
        for item in payload["education"]:
            for label, value in [
                ("毕业院校", item.get("school", "")),
                ("专业", item.get("major", "")),
                ("学历", item.get("degree", "")),
                ("时间", date_range(item)),
            ]:
                if value:
                    education_rows.append((label, value))
        add_basic_info(document, education_rows)

    certificates = payload["certificates"]
    if certificates:
        add_section_heading(document, "证书与奖项")
        add_arrow_bullets(document, certificates, size=9)

    if payload["summary"] and not payload["projects"] and not payload["work_experience"]:
        add_section_heading(document, "自我评价")
        add_arrow_bullets(document, [payload["summary"]], size=9)

    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "resume.docx"
    document.save(docx_path)

    (out_dir / "resume.txt").write_text(render_plain_text(facts, selected_template), encoding="utf-8")
    (out_dir / "resume-render.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    metadata = {
        "selected_template": selected_template,
        "section_order": ["基本信息", "求职意向", "开发技能", "工作经历", "项目经历", "教育背景", "自我评价"],
        "style_source": STYLE_SOURCE,
        "style_tokens": {"blue": BLUE, "gold": GOLD},
    }
    (out_dir / "metadata.json").write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")
    return docx_path


def assert_style_tokens(docx_path: Path) -> None:
    with zipfile.ZipFile(docx_path) as archive:
        xml = archive.read("word/document.xml")
    missing = [token for token in [BLUE.encode(), GOLD.encode()] if token not in xml]
    if missing:
        raise RuntimeError(f"generated DOCX is missing canonical style tokens: {missing}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--facts", required=True, type=Path)
    parser.add_argument("--template-index", required=True, type=Path)
    parser.add_argument("--out-dir", required=True, type=Path)
    parser.add_argument("--language", default="zh-CN")
    args = parser.parse_args()

    facts = load_json(args.facts)
    index = load_json(args.template_index)
    selected_template = select_template(index)
    docx_path = render_resume(facts, selected_template, args.out_dir)
    assert_style_tokens(docx_path)
    print(f"generated {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
